from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parent.parent
DOCS = ROOT / "docs"

INPUT_FILES = [
    DOCS / "production_readiness_report.md",
    DOCS / "presentation_demo_script.md",
]
OUTPUT_FILE = DOCS / "PCOSense_AppV3_Submission_Draft.docx"


def add_markdown_as_plain_sections(document: Document, path: Path) -> None:
    document.add_heading(path.stem.replace("_", " ").title(), level=1)
    for raw_line in path.read_text(encoding="utf-8").splitlines():
        line = raw_line.rstrip()
        if not line:
            document.add_paragraph("")
            continue
        if line.startswith("# "):
            document.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            document.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            document.add_heading(line[4:].strip(), level=3)
        elif line.startswith("- "):
            document.add_paragraph(line[2:].strip(), style="List Bullet")
        elif line.startswith("1. "):
            document.add_paragraph(line[3:].strip(), style="List Number")
        else:
            document.add_paragraph(line)


def main() -> None:
    doc = Document()
    doc.add_heading("PCOSense App V3 Submission Draft", level=0)
    doc.add_paragraph("Generated from project docs for final review and submission packaging.")

    for src in INPUT_FILES:
        if src.exists():
            add_markdown_as_plain_sections(doc, src)
        else:
            doc.add_heading(src.name, level=1)
            doc.add_paragraph("File not found.")

    doc.save(OUTPUT_FILE)
    print(f"Created: {OUTPUT_FILE}")


if __name__ == "__main__":
    main()
